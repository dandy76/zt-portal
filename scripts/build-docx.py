#!/usr/bin/env python3
"""
Convert USER_GUIDE.md to USER_GUIDE.docx with embedded screenshots.
Usage: python3 scripts/build-docx.py
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
MD_FILE = ROOT / "USER_GUIDE.md"
OUT_FILE = ROOT / "USER_GUIDE.docx"
SCREENSHOTS = ROOT

IMG_RE = re.compile(r'!\[(.*?)\]\((.*?)\)')
CODE_FENCE = re.compile(r'^```(\w+)?$')
HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)$')
BOLD_ITALIC_RE = re.compile(r'(\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`)')
LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def add_formatted_runs(paragraph, text):
    """Parse a line with **bold**, *italic*, `code`, [link](url) and add as runs."""
    text = LINK_RE.sub(r'\1', text)
    pos = 0
    for m in BOLD_ITALIC_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(2):
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(3):
            r = paragraph.add_run(m.group(3))
            r.italic = True
        elif m.group(4):
            r = paragraph.add_run(m.group(4))
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x40, 0x80)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build():
    if not MD_FILE.exists():
        print(f"ERROR: {MD_FILE} not found", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    in_code = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        header = table_rows[0]
        body = table_rows[2:] if len(table_rows) > 1 else []
        t = doc.add_table(rows=1 + len(body), cols=len(header))
        t.style = 'Light Grid Accent 1'
        for i, h in enumerate(header):
            cell = t.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
        for r_idx, row in enumerate(body, 1):
            for c_idx, cell_text in enumerate(row):
                if c_idx < len(t.rows[r_idx].cells):
                    t.rows[r_idx].cells[c_idx].text = cell_text
        doc.add_paragraph()
        table_rows = []

    lines = MD_FILE.read_text(encoding='utf-8').splitlines()
    for line in lines:
        fence = CODE_FENCE.match(line)
        if fence:
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run('\n'.join(code_lines))
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        # Table handling
        if line.startswith('|') and line.strip().endswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            flush_table()
            in_table = False

        # Images
        img = IMG_RE.search(line)
        if img:
            alt, path = img.group(1), img.group(2)
            img_path = ROOT / path
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(6.0))
                    last = doc.paragraphs[-1]
                    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if alt:
                        cap = doc.add_paragraph(alt)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cap.runs:
                            run.italic = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                except Exception as e:
                    doc.add_paragraph(f"[Image not rendered: {path} - {e}]")
            else:
                doc.add_paragraph(f"[Missing image: {path}]")
            continue

        # Headings
        h = HEADING_RE.match(line)
        if h:
            level = len(h.group(1))
            doc.add_heading(h.group(2), level=min(level, 4))
            continue

        # Horizontal rule
        if line.strip() in ('---', '***'):
            doc.add_paragraph('_' * 80).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            continue

        # Blockquote
        if line.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_formatted_runs(p, line.lstrip('> ').strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            continue

        # Bullets
        if line.lstrip().startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, line.lstrip()[2:])
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line.lstrip()):
            p = doc.add_paragraph(style='List Number')
            add_formatted_runs(p, re.sub(r'^\d+\.\s+', '', line.lstrip()))
            continue

        # Empty
        if not line.strip():
            doc.add_paragraph()
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, line)

    if in_table:
        flush_table()

    doc.save(OUT_FILE)
    print(f"Created: {OUT_FILE} ({OUT_FILE.stat().st_size // 1024} KB)")


if __name__ == '__main__':
    build()
